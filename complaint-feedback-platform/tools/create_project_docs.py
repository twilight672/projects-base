from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("output/doc")
OUT_DIR.mkdir(parents=True, exist_ok=True)


BLUE = "1D5F93"
LIGHT_BLUE = "EAF2F8"
LINE = "D9E2EC"
DARK = "1E2A3A"
MUTED = "66758A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_layout(table, widths: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11.5, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(120)
    r = cover.add_run(title)
    set_run_font(r, 24, True, BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, 12, False, MUTED)

    rows = [
        ("项目名称", "民生诉求反馈平台"),
        ("文档版本", "V1.0"),
        ("编制日期", "2026年4月25日"),
        ("适用范围", "系统开发、交付验收、部署上线与后续维护交接"),
    ]
    doc.add_paragraph()
    for k, v in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rk = p.add_run(f"{k}：")
        rv = p.add_run(v)
        set_run_font(rk, 10.5, True, BLUE)
        set_run_font(rv, 10.5, False, DARK)

    doc.add_page_break()
    return doc


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    for row_data in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(5)
        for i, value in enumerate(row_data):
            key = headers[i]
            rk = p.add_run(f"{key}：")
            rv = p.add_run(value)
            set_run_font(rk, 10, True, BLUE)
            set_run_font(rv, 10, False, DARK)
            if i < len(row_data) - 1:
                sep = p.add_run("；")
                set_run_font(sep, 10, False, MUTED)
    doc.add_paragraph()


def add_note(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, 10.5, True, BLUE)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(body)
    for run in p.runs:
        set_run_font(run, 10)


def build_technical_doc() -> Path:
    doc = setup_doc("民生诉求反馈平台技术文档", "系统架构、功能模块、接口、数据与测试说明")

    doc.add_heading("一、项目概述", level=1)
    doc.add_paragraph(
        "民生诉求反馈平台用于承接公众投诉、问题反馈和办理进度查询，支持手机扫码和电脑端访问。"
        "系统包含公众提交入口、市县两级管理后台、附件上传、办理闭环、统计导出和移动端适配能力。"
    )
    add_bullets(
        doc,
        [
            "公众端支持匿名与实名两种提交模式，实名信息仅按权限受控展示。",
            "县级账号只处理本县区事项，市级账号用于全市总览、统计、督办和导出。",
            "系统采用单 Node 服务部署模式，生产环境由 Express 同时托管前端静态资源和后端 API。",
            "数据使用 sql.js 持久化 SQLite 文件，附件上传到本地 uploads 目录，便于中小规模部署和交付。"
        ],
    )

    doc.add_heading("二、技术选型", level=1)
    add_table(
        doc,
        ["层级", "技术/组件", "用途说明"],
        [
            ["前端", "Vite + React + TypeScript", "构建公众端、查询页和管理后台，支持组件化和类型检查。"],
            ["UI", "CSS + lucide-react", "行政简约风界面、图标按钮、移动端响应式布局。"],
            ["后端", "Node.js + Express + TypeScript", "提供登录、投诉、附件、后台管理和导出接口。"],
            ["数据库", "sql.js + SQLite 文件", "数据持久化到 data/complaints.sqlite，避免原生数据库依赖。"],
            ["附件", "multer", "支持图片、PDF、Word、Excel 上传，默认单文件 10MB、最多 5 个。"],
            ["二维码", "qrcode", "根据当前访问地址生成扫码入口。"],
            ["测试", "TypeScript + Playwright", "覆盖类型检查、构建、接口 smoke 和电脑/手机端流程测试。"],
        ],
        [2.2, 4.0, 9.0],
    )

    doc.add_heading("三、系统架构", level=1)
    doc.add_paragraph(
        "系统采用前后端同仓库开发、生产环境单服务托管的架构。开发时前端由 Vite 提供热更新，"
        "后端由 tsx 运行 TypeScript 服务；上线时先构建 dist/client 和 dist/server，再由 Express 提供静态页面与 API。"
    )
    add_table(
        doc,
        ["目录/文件", "说明"],
        [
            ["client/", "前端 React 应用，包含页面、样式和 Vite 配置。"],
            ["server/", "后端 Express 服务，包含接口、数据库、登录、附件上传和 smoke 检查。"],
            ["data/", "SQLite 持久化文件目录，生产环境需要备份和保留。"],
            ["uploads/", "用户附件上传目录，生产环境需要备份和保留。"],
            ["dist/", "npm run build 后生成的生产构建产物。"],
            ["tests/", "Playwright 端到端测试用例。"],
            ["package.json", "统一维护安装依赖、开发、构建、启动和测试脚本。"],
        ],
        [4.0, 11.0],
    )

    doc.add_heading("四、功能模块", level=1)
    add_table(
        doc,
        ["模块", "主要功能", "当前实现状态"],
        [
            ["公众反馈", "匿名/实名提交、县区选择、类型选择、标题、内容、附件上传、提交编号。", "已实现"],
            ["进度查询", "按编号查询事项状态、办理日志、附件、补充入口和评价入口。", "已实现"],
            ["补充材料", "待补充状态下，公众可补充说明和附件。", "已实现"],
            ["满意度评价", "事项办结后，公众可提交 1-5 分评价和意见。", "已实现"],
            ["县级后台", "按县区隔离数据，支持受理、转派、退回补充、办结。", "已实现"],
            ["市级后台", "全市统计、县区排行、脱敏查看、督办、CSV 导出。", "已实现"],
            ["附件管理", "后台详情可查看附件列表并打开文件。", "已实现"],
        ],
        [2.7, 9.0, 2.5],
    )

    doc.add_heading("五、账号与权限", level=1)
    add_table(
        doc,
        ["账号", "密码", "角色", "权限范围"],
        [
            ["city", "123456", "市级管理员", "查看全市统计和事项，手机号脱敏，可督办和导出。"],
            ["county-east", "123456", "东城区办理员", "仅查看和办理东城区事项，可见完整实名联系方式。"],
            ["county-west", "123456", "西城区办理员", "仅查看和办理西城区事项，可见完整实名联系方式。"],
        ],
        [3.2, 2.2, 3.2, 6.5],
    )
    add_note(
        doc,
        "隐私规则",
        "实名投诉的姓名和电话不向公众查询接口返回。后台中，县级办理员因联系核实需要可见完整电话；市级账号默认只显示脱敏电话。"
    )

    doc.add_heading("六、核心接口", level=1)
    add_table(
        doc,
        ["接口", "方法", "说明"],
        [
            ["/api/auth/login", "POST", "账号密码登录，写入 HttpOnly cookie 登录态。"],
            ["/api/auth/me", "GET", "获取当前登录用户信息。"],
            ["/api/complaints", "POST", "公众提交投诉，支持 multipart/form-data 附件。"],
            ["/api/complaints/:code", "GET", "公众按编号查询进度，不返回实名联系方式。"],
            ["/api/complaints/:code/supplement", "POST", "公众提交补充说明和附件。"],
            ["/api/complaints/:code/evaluation", "POST", "事项办结后提交满意度评价。"],
            ["/api/admin/stats", "GET", "后台统计，自动按市级/县级权限过滤。"],
            ["/api/admin/complaints", "GET", "后台列表，支持状态、县区、关键词筛选。"],
            ["/api/admin/complaints/:id/actions", "POST", "办理动作：受理、转派、退回补充、办结、督办。"],
            ["/api/admin/export.csv", "GET", "按当前权限导出 CSV。"],
        ],
        [5.0, 2.0, 8.0],
    )

    doc.add_heading("七、数据设计", level=1)
    add_table(
        doc,
        ["数据表", "用途", "关键字段"],
        [
            ["users", "后台账号与权限", "username、password_hash、role、county、label"],
            ["complaints", "投诉主表", "code、mode、county、category、title、content、status、assignee、evaluation"],
            ["attachments", "附件表", "complaint_id、kind、original_name、stored_name、mime_type、size、path"],
            ["logs", "办理日志", "complaint_id、action、actor、note、created_at"],
        ],
        [3.0, 4.0, 8.0],
    )

    doc.add_heading("八、安全与运维注意事项", level=1)
    add_bullets(
        doc,
        [
            "生产环境应修改默认密码，并优先通过 HTTPS 访问系统。",
            "data/complaints.sqlite 和 uploads/ 是核心业务资产，必须定期备份。",
            "当前版本适合中小规模反馈系统，若后续并发和数据量明显增大，建议迁移到 PostgreSQL 或 MySQL。",
            "公网部署时建议使用 Nginx 反向代理，并限制服务器防火墙只开放 80/443 端口。",
            "附件目录应限制可执行文件上传，当前实现仅允许图片、PDF、Word、Excel 类型。"
        ],
    )

    doc.add_heading("九、测试与验收", level=1)
    add_table(
        doc,
        ["检查项", "命令/方式", "验收标准"],
        [
            ["类型检查", "npm run typecheck", "前端和后端 TypeScript 无类型错误。"],
            ["构建检查", "npm run build", "生成 dist/client 和 dist/server。"],
            ["接口 smoke", "npm run smoke", "健康检查、提交、查询接口可用。"],
            ["端到端测试", "npm run test:e2e", "公众提交、查询、县级办理、市级督办、移动端适配通过。"],
            ["总检查", "npm run check", "类型检查、构建和 smoke 全部通过。"],
        ],
        [3.2, 4.0, 7.8],
    )

    doc.add_heading("十、交付清单", level=1)
    add_bullets(
        doc,
        [
            "完整源代码：client、server、tests、配置文件和 package-lock.json。",
            "生产构建产物：dist/client、dist/server。",
            "业务数据文件：data/complaints.sqlite。",
            "附件目录：uploads。",
            "技术文档与部署教程：本次生成的两个 Word 文档。"
        ],
    )

    path = OUT_DIR / "民生诉求反馈平台_技术文档.docx"
    doc.save(path)
    return path


def build_deploy_doc() -> Path:
    doc = setup_doc("民生诉求反馈平台部署教程", "服务器上线、反向代理、备份、更新与故障处理")

    doc.add_heading("一、部署目标", level=1)
    doc.add_paragraph(
        "本教程用于将民生诉求反馈平台部署到公网服务器。部署完成后，用户可通过域名或服务器 IP 访问，"
        "手机扫码进入公众反馈页面，管理员通过后台账号登录处理事项。"
    )
    add_note(
        doc,
        "推荐方式",
        "推荐使用 Linux 服务器 + Node.js 20+ + PM2 + Nginx + HTTPS。若暂时没有域名，也可以先使用 http://服务器IP:3000 进行验收。"
    )

    doc.add_heading("二、服务器准备", level=1)
    add_table(
        doc,
        ["项目", "最低要求", "建议配置"],
        [
            ["操作系统", "Windows Server 或 Linux", "Ubuntu 22.04 / Debian 12 / CentOS Stream"],
            ["Node.js", "20+", "20 LTS 或 22 LTS"],
            ["内存", "1GB", "2GB 以上"],
            ["磁盘", "5GB 可用空间", "20GB 以上，便于附件和备份"],
            ["端口", "3000 或 80/443", "公网仅开放 80/443，3000 走内网反代"],
        ],
        [3.0, 4.3, 7.5],
    )

    doc.add_heading("三、上传项目", level=1)
    doc.add_paragraph("将本机项目目录上传到服务器，例如：")
    add_note(doc, "本机项目目录", r"C:\Users\33072\Desktop\index")
    doc.add_paragraph("服务器建议目录：")
    add_note(doc, "Linux 示例", "/opt/complaint-feedback-platform")
    add_note(doc, "Windows Server 示例", r"D:\sites\complaint-feedback-platform")

    doc.add_heading("四、Linux 部署步骤", level=1)
    add_numbered(
        doc,
        [
            "安装 Node.js 20+，确认 node -v 和 npm -v 可用。",
            "进入项目目录：cd /opt/complaint-feedback-platform。",
            "安装依赖：npm install。",
            "执行检查：npm run check。",
            "构建生产包：npm run build。",
            "启动服务：npm run start。",
            "浏览器访问：http://服务器IP:3000，确认页面和后台可用。"
        ],
    )
    add_note(
        doc,
        "后台账号",
        "市级 city / 123456；县级 county-east / 123456；县级 county-west / 123456。上线后建议立即修改默认密码。"
    )

    doc.add_heading("五、PM2 守护进程部署", level=1)
    doc.add_paragraph("生产环境建议使用 PM2 保持服务后台运行，并支持开机自启。")
    add_table(
        doc,
        ["步骤", "命令"],
        [
            ["安装 PM2", "npm install -g pm2"],
            ["启动项目", "pm2 start dist/server/index.js --name complaint-platform"],
            ["查看状态", "pm2 status"],
            ["查看日志", "pm2 logs complaint-platform"],
            ["保存进程", "pm2 save"],
            ["设置开机自启", "pm2 startup"],
            ["重启服务", "pm2 restart complaint-platform"],
        ],
        [4.0, 10.8],
    )

    doc.add_heading("六、Nginx 反向代理", level=1)
    doc.add_paragraph("若使用域名访问，建议 Nginx 监听 80/443，并代理到本机 3000 端口。")
    add_note(
        doc,
        "Nginx 配置示例",
        "server {\n"
        "    listen 80;\n"
        "    server_name your-domain.com;\n\n"
        "    client_max_body_size 50m;\n\n"
        "    location / {\n"
        "        proxy_pass http://127.0.0.1:3000;\n"
        "        proxy_set_header Host $host;\n"
        "        proxy_set_header X-Real-IP $remote_addr;\n"
        "        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n"
        "        proxy_set_header X-Forwarded-Proto $scheme;\n"
        "    }\n"
        "}"
    )
    add_bullets(
        doc,
        [
            "配置完成后执行 nginx -t 检查语法。",
            "语法正常后执行 systemctl reload nginx 使配置生效。",
            "HTTPS 可使用服务器面板或 Certbot 申请证书。"
        ],
    )

    doc.add_heading("七、Windows Server 部署步骤", level=1)
    add_numbered(
        doc,
        [
            "安装 Node.js 20+，打开 PowerShell 确认 node -v 和 npm -v。",
            "将项目上传到 D:\\sites\\complaint-feedback-platform。",
            "进入目录后执行 npm install。",
            "执行 npm run check 验证项目。",
            "执行 npm run build 生成生产产物。",
            "执行 npm run start 启动服务。",
            "如需长期运行，可使用 PM2 或 NSSM 将 node dist/server/index.js 注册为服务。"
        ],
    )

    doc.add_heading("八、数据与附件备份", level=1)
    add_table(
        doc,
        ["资产", "路径", "备份建议"],
        [
            ["数据库", "data/complaints.sqlite", "每天备份一次，更新发布前必须备份。"],
            ["附件", "uploads/", "每天增量备份，保留最近 30 天。"],
            ["配置与代码", "项目根目录", "每次发布保留上一版本压缩包。"],
        ],
        [3.0, 5.0, 7.0],
    )
    add_note(
        doc,
        "重要提醒",
        "更新代码时不要删除 data/ 和 uploads/。这两个目录包含真实业务数据和附件，误删会造成数据丢失。"
    )

    doc.add_heading("九、更新发布流程", level=1)
    add_numbered(
        doc,
        [
            "停止服务或保持旧服务运行，先备份 data/ 和 uploads/。",
            "上传新代码，保留原 data/ 和 uploads/。",
            "执行 npm install，安装新增依赖。",
            "执行 npm run check，确认构建和接口检查通过。",
            "执行 pm2 restart complaint-platform 或重新运行 npm run start。",
            "打开首页、提交页、查询页和后台，完成一次冒烟验收。"
        ],
    )

    doc.add_heading("十、上线验收清单", level=1)
    add_table(
        doc,
        ["验收项", "检查方法", "通过标准"],
        [
            ["首页访问", "打开域名或 IP", "页面正常加载，无空白和明显错位。"],
            ["扫码访问", "手机扫描页面二维码", "手机端能打开同一系统地址。"],
            ["匿名提交", "填写匿名反馈", "生成 MS 开头查询编号。"],
            ["实名提交", "填写姓名和电话", "县级后台可见完整电话，市级脱敏。"],
            ["附件上传", "上传图片/PDF/Word/Excel", "后台详情可查看附件列表。"],
            ["县级办理", "county-east 登录", "仅看到东城区事项，可受理、转派、退回补充、办结。"],
            ["市级督办", "city 登录", "可查看全市统计、督办、导出 CSV。"],
            ["备份检查", "查看 data 和 uploads", "目录存在且有写入权限。"],
        ],
        [3.2, 5.0, 6.8],
    )

    doc.add_heading("十一、常见问题处理", level=1)
    add_table(
        doc,
        ["问题", "可能原因", "处理方式"],
        [
            ["页面打不开", "服务未启动或端口未开放", "执行 pm2 status / netstat，检查 3000、80、443 端口。"],
            ["接口 404", "未执行 npm run build 或 dist 缺失", "重新执行 npm run build，再重启服务。"],
            ["附件无法上传", "uploads 无写入权限或 Nginx 限制过小", "检查目录权限，并设置 client_max_body_size 50m。"],
            ["数据丢失", "发布时覆盖了 data 目录", "从最近备份恢复 data/complaints.sqlite。"],
            ["手机扫码访问失败", "二维码指向内网或未绑定公网域名", "使用公网域名访问页面后重新扫码。"],
            ["npm 安装较慢", "官方源网络不稳定", "可临时使用 npm install --registry=https://registry.npmmirror.com。"],
        ],
        [3.5, 5.0, 6.5],
    )

    doc.add_heading("十二、生产安全建议", level=1)
    add_bullets(
        doc,
        [
            "上线后立即修改默认账号密码。",
            "域名上线后启用 HTTPS，避免明文传输实名信息。",
            "限制服务器登录权限，定期更新系统补丁。",
            "对 data/ 和 uploads/ 设置定期异地备份。",
            "如后续数据规模增大，建议升级为 MySQL/PostgreSQL 并增加操作审计。"
        ],
    )

    path = OUT_DIR / "民生诉求反馈平台_部署教程.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    technical = build_technical_doc()
    deployment = build_deploy_doc()
    print(technical)
    print(deployment)
